import os
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('RAG Document Assistant - v4 Pipeline Documentation', 0)

# Section 1
doc.add_heading('1. Document Upload and Indexing', level=1)
doc.add_paragraph('Trigger: User uploads PDF/DOCX/TXT and clicks "Process and Index"')
doc.add_paragraph('')

t = doc.add_table(rows=8, cols=3)
t.style = 'Table Grid'
headers = [('Step', 'Action', 'Details')]
data = [
    ('1', 'Clear previous data', 'All old documents, FAISS index, chunks.pkl, and semantic cache are wiped'),
    ('2', 'Save file', 'Uploaded file saved to uploads/ with a UUID filename'),
    ('3', 'Load document', 'load_document() reads PDF/DOCX/TXT into raw text'),
    ('4', 'Split into chunks', 'split_documents() splits text into ~800 char chunks with 150 char overlap, adding metadata (section, company, role, etc.)'),
    ('5', 'FAISS indexing', 'Chunks embedded using BAAI/bge-base-en-v1.5 (768-dim) and stored in FAISS vector index'),
    ('6', 'BM25 store', 'Raw chunks pickled to chunks.pkl for keyword search'),
    ('7', 'Registry update', 'documents.json updated with doc ID, filename, chunk count'),
]
t.cell(0, 0).text = 'Step'
t.cell(0, 1).text = 'Action'
t.cell(0, 2).text = 'Details'
for i, (s, a, d) in enumerate(data, 1):
    t.cell(i, 0).text = s
    t.cell(i, 1).text = a
    t.cell(i, 2).text = d

doc.add_paragraph('')
doc.add_paragraph(
    'Key Design Decision: Only ONE document is active at a time. '
    'Uploading a new document completely replaces the previous one (vector store + cache fully cleared first).',
    style='Intense Quote'
)

# Section 2
doc.add_heading('2. Query Processing', level=1)

doc.add_heading('Step 0a: Conversation-Aware Query Rewrite', level=2)
doc.add_paragraph(
    'Checks if there is prior conversation AND the question contains context words '
    '(next, previous, that, this, it, second, third, last, etc.) or is 3 words or fewer.'
)
doc.add_paragraph('If YES: LLM rewrites the question into a standalone broad query.')
doc.add_paragraph(
    'Example: "next company?" becomes "What are all the companies worked at in chronological order? '
    'Specifically the second one."'
)
doc.add_paragraph(
    'The rewriter does NOT include specific previous answers in the query to avoid biasing retrieval.'
)

doc.add_heading('Step 0b: Semantic Cache Check (skipped for follow-ups)', level=2)
doc.add_paragraph(
    'If question was NOT rewritten: embed question, compare against cached Q&A pairs (cosine similarity).'
)
doc.add_paragraph('If similarity >= 0.55: ask validator LLM "is this the same question?"')
doc.add_paragraph('Validator checks SCOPE (filters added/removed?) and INTENT (list vs count?).')
doc.add_paragraph('If YES: return cached answer instantly. "Not found" answers are never cached.')

doc.add_heading('Step 1: Semantic Search (FAISS)', level=2)
doc.add_paragraph(
    'Expand query with topic keywords (e.g. "skills" adds "technical skills programming languages '
    'tools technologies expertise"). FAISS similarity search returns top 15 candidates.'
)

doc.add_heading('Step 2: Keyword Search (BM25)', level=2)
doc.add_paragraph(
    'BM25 retriever (term frequency-based) over all chunks returns top 15 candidates. '
    'Catches exact keyword matches that embedding search might miss.'
)

doc.add_heading('Step 3: Merge and Deduplicate', level=2)
doc.add_paragraph('Combine FAISS + BM25 results. Remove duplicates by comparing first 200 characters.')

doc.add_heading('Step 4: Cross-Encoder Reranking + Section Affinity Boost', level=2)
doc.add_paragraph(
    'Cross-encoder (ms-marco-MiniLM-L-6-v2) scores each (question, chunk) pair for relevance.'
)
doc.add_paragraph(
    'Section affinity boost: +2.5 score added when question topic aligns with chunk section metadata '
    '(e.g. question about "experience" boosts experience chunks).'
)
doc.add_paragraph('Keep top chunks within 3.0 points of best score (maximum 5 chunks).')

doc.add_heading('Step 5: LLM Generation (Google Gemini 3.1 Flash-Lite)', level=2)
doc.add_paragraph(
    'Build context from top chunks with metadata labels (Company, Role, Period, Section). '
    'Include conversation history (last 3 Q&A pairs) for follow-ups.'
)
doc.add_paragraph('Prompt rules enforce:')
rules = [
    'Answer ONLY from retrieved passages',
    'Say "not mentioned" if info is not there',
    'Do not repeat previously given answers',
    'Indicate "(currently working)" for present employment',
    'List ALL items found (do not summarize)',
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('Step 6: Format Sources', level=2)
doc.add_paragraph('Attach chunk snippets + source file + page number for transparency.')

doc.add_heading('Step 7: Cache Result', level=2)
doc.add_paragraph(
    'Store question embedding + answer in semantic cache. Skip if answer is "not mentioned in the document".'
)

# Section 3
doc.add_heading('3. Conversation Memory', level=1)
t2 = doc.add_table(rows=7, cols=2)
t2.style = 'Table Grid'
mem_data = [
    ('Feature', 'Implementation'),
    ('Storage', 'All Q&A pairs in Streamlit session state'),
    ('Rewrite trigger', 'Context words or <= 3 word questions'),
    ('History window', 'Last 3 Q&A pairs (6 messages)'),
    ('LLM context', 'Conversation history injected into RAG prompt'),
    ('Cache bypass', 'Rewritten questions always skip cache'),
    ('Clear', '"Clear Chat History" button resets conversation'),
]
for i, (f, imp) in enumerate(mem_data):
    t2.cell(i, 0).text = f
    t2.cell(i, 1).text = imp

# Section 4
doc.add_heading('4. Architecture Components', level=1)
t3 = doc.add_table(rows=9, cols=3)
t3.style = 'Table Grid'
arch_data = [
    ('Component', 'Technology', 'Purpose'),
    ('Frontend', 'Streamlit', 'Chat UI, file upload, document management'),
    ('Embeddings', 'BAAI/bge-base-en-v1.5', '768-dim dense vectors for semantic search'),
    ('Vector Store', 'FAISS', 'Fast similarity search over embeddings'),
    ('Keyword Search', 'BM25Retriever', 'Term-frequency based retrieval'),
    ('Reranker', 'cross-encoder/ms-marco-MiniLM-L-6-v2', 'Precise relevance scoring'),
    ('LLM', 'Google Gemini 3.1 Flash-Lite', 'Answer generation + cache validation + query rewriting'),
    ('Cache', 'Custom semantic cache (cosine + LLM validator)', 'Avoid redundant LLM calls'),
    ('Persistence', 'FAISS index + pickle files', 'Survive app restarts'),
]
for i, (c, tech, p) in enumerate(arch_data):
    t3.cell(i, 0).text = c
    t3.cell(i, 1).text = tech
    t3.cell(i, 2).text = p

# Section 5
doc.add_heading('5. Configuration (config.py)', level=1)
t4 = doc.add_table(rows=13, cols=3)
t4.style = 'Table Grid'
cfg_data = [
    ('Parameter', 'Value', 'Description'),
    ('EMBEDDING_MODEL', 'BAAI/bge-base-en-v1.5', 'Embedding model (768-dim)'),
    ('CROSS_ENCODER_MODEL', 'ms-marco-MiniLM-L-6-v2', 'Reranker model'),
    ('LLM_MODEL', 'gemini-3.1-flash-lite', 'Generation model'),
    ('LLM_TEMPERATURE', '0.1', 'Low temperature for precise answers'),
    ('CHUNK_SIZE', '800', 'Characters per chunk'),
    ('CHUNK_OVERLAP', '150', 'Overlap between chunks'),
    ('INITIAL_RETRIEVAL_K', '15', 'Candidates from each retriever'),
    ('MAX_RETRIEVAL_DOCS', '5', 'Top chunks sent to LLM'),
    ('SCORE_GAP', '3.0', 'Max score distance from best chunk'),
    ('CACHE_CANDIDATE_FLOOR', '0.55', 'Min cosine sim to consider cache hit'),
    ('CACHE_TTL_DAYS', '7', 'Cache entry lifetime'),
    ('CACHE_MAX_SIZE', '500', 'Max cached entries (LRU eviction)'),
]
for i, (param, val, desc) in enumerate(cfg_data):
    t4.cell(i, 0).text = param
    t4.cell(i, 1).text = val
    t4.cell(i, 2).text = desc

# Section 6
doc.add_heading('6. Key v4 Improvements', level=1)
improvements = [
    'Single-document replacement: Uploading a new doc fully clears the old one (no stale data)',
    'Conversation memory: Follow-up questions understand context from prior Q&A',
    'Smart cache bypass: Conversation-dependent questions skip cache to avoid wrong matches',
    'No caching of "not found": Prevents negative answers from being served to rephrased questions',
    'Current employment indicator: Answers include "(currently working)" when applicable',
    'Broad rewrite strategy: Follow-up queries rewritten to find ALL items, preventing retrieval bias',
]
for imp in improvements:
    doc.add_paragraph(imp, style='List Number')

# Save
output_path = r'D:\OneDrive - HealthEdge Software, Inc\Documents\Personal\rag-doc-assistant-v3\RAG_Pipeline_Documentation_v4.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
